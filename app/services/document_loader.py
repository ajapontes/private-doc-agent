"""
Local document loader service.

This module lists supported local documents and extracts their textual
content using the appropriate reader for each file type.

Supported formats:
- TXT and Markdown files encoded as UTF-8.
- PDF files containing extractable text.
- Microsoft Word DOCX documents.

Scanned PDFs that contain only images are not supported because this
service does not perform optical character recognition (OCR).

Logging strategy:
- Logs operational events such as directory scanning and document reading.
- Logs metadata such as filename, extension, file size, page count,
  paragraph count, and extracted character count.
- Does not log document content to avoid exposing private information.
"""

import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.config import INPUT_DIR, SUPPORTED_EXTENSIONS


logger = logging.getLogger(__name__)


def list_documents() -> list[dict]:
    """
    Lists supported documents from the local input folder.

    Returns:
        A list of dictionaries containing the filename, extension, size in
        bytes, and relative path of each supported document.
    """
    logger.info(
        "Scanning input directory for supported documents. "
        "input_dir=%s supported_extensions=%s",
        INPUT_DIR,
        sorted(SUPPORTED_EXTENSIONS),
    )

    INPUT_DIR.mkdir(parents=True, exist_ok=True)

    documents = []

    for file_path in INPUT_DIR.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            document_metadata = {
                "filename": file_path.name,
                "extension": file_path.suffix.lower(),
                "size_bytes": file_path.stat().st_size,
                "path": str(file_path.relative_to(INPUT_DIR.parent.parent)),
            }

            documents.append(document_metadata)

            logger.info(
                "Supported document found. "
                "filename=%s extension=%s size_bytes=%s",
                document_metadata["filename"],
                document_metadata["extension"],
                document_metadata["size_bytes"],
            )

    logger.info("Document scan completed. count=%s", len(documents))

    return documents


def _read_text_document(file_path: Path) -> str:
    """
    Reads a UTF-8 plain-text or Markdown document.

    Args:
        file_path: Path of the document to read.

    Returns:
        Extracted document text.

    Raises:
        ValueError: If the document cannot be decoded as UTF-8.
    """
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError as error:
        raise ValueError(
            f"Document is not valid UTF-8 text: {file_path.name}"
        ) from error


def _read_pdf_document(file_path: Path) -> str:
    """
    Extracts text from every page of a PDF document.

    Args:
        file_path: Path of the PDF document.

    Returns:
        Extracted text separated by page.

    Raises:
        ValueError: If the PDF is damaged, encrypted, or contains no
            extractable text.
    """
    try:
        reader = PdfReader(file_path)

        if reader.is_encrypted:
            raise ValueError(f"Encrypted PDF documents are not supported: {file_path.name}")

        page_texts = []

        for page in reader.pages:
            extracted_text = page.extract_text()

            if extracted_text and extracted_text.strip():
                page_texts.append(extracted_text.strip())

        content = "\n\n".join(page_texts)

        logger.info(
            "PDF extraction completed. filename=%s pages=%s pages_with_text=%s",
            file_path.name,
            len(reader.pages),
            len(page_texts),
        )
    except ValueError:
        raise
    except Exception as error:
        raise ValueError(
            f"Unable to read PDF document: {file_path.name}"
        ) from error

    if not content.strip():
        raise ValueError(
            f"PDF document contains no extractable text: {file_path.name}"
        )

    return content


def _read_docx_document(file_path: Path) -> str:
    """
    Extracts text from the paragraphs of a DOCX document.

    Args:
        file_path: Path of the DOCX document.

    Returns:
        Extracted paragraphs separated by line breaks.

    Raises:
        ValueError: If the document is damaged or contains no extractable text.
    """
    try:
        document = Document(file_path)
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        content = "\n\n".join(paragraphs)

        logger.info(
            "DOCX extraction completed. filename=%s paragraphs=%s paragraphs_with_text=%s",
            file_path.name,
            len(document.paragraphs),
            len(paragraphs),
        )
    except Exception as error:
        raise ValueError(
            f"Unable to read DOCX document: {file_path.name}"
        ) from error

    if not content.strip():
        raise ValueError(
            f"DOCX document contains no extractable text: {file_path.name}"
        )

    return content


def read_document(filename: str) -> str:
    """
    Reads a supported local document by filename.

    The reader is selected from the document extension.

    Args:
        filename: Name of the document located in the configured input folder.

    Returns:
        Extracted document content as plain text.

    Raises:
        FileNotFoundError: If the document does not exist.
        ValueError: If the extension is unsupported or the document cannot be
            processed.
    """
    logger.info("Reading document requested. filename=%s", filename)

    file_path = INPUT_DIR / filename
    extension = file_path.suffix.lower()

    if not file_path.exists() or not file_path.is_file():
        logger.warning("Document not found. filename=%s path=%s", filename, file_path)
        raise FileNotFoundError(f"Document not found: {filename}")

    if extension not in SUPPORTED_EXTENSIONS:
        logger.warning(
            "Unsupported document extension. filename=%s extension=%s",
            filename,
            extension,
        )
        raise ValueError(f"Unsupported file extension: {file_path.suffix}")

    readers = {
        ".txt": _read_text_document,
        ".md": _read_text_document,
        ".pdf": _read_pdf_document,
        ".docx": _read_docx_document,
    }

    reader = readers.get(extension)

    if reader is None:
        raise ValueError(f"No document reader configured for extension: {extension}")

    content = reader(file_path)

    logger.info(
        "Document read completed. filename=%s extension=%s size_chars=%s",
        filename,
        extension,
        len(content),
    )

    return content